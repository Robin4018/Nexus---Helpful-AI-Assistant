import pypdf
import docx
import os
import json
from google.genai import types

def extract_text_from_file(file_path, file_type):
    """
    Given a local file path and a mime-type, extract readable text if it is a document.
    Supports PDF, DOCX (including tables), CSV, Markdown, text, and other files.
    """
    text_content = ""
    try:
        ft = file_type.lower()
        file_path_lower = file_path.lower()
        
        # 1. PDF Documents
        if 'pdf' in ft or file_path_lower.endswith('.pdf'):
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for i, page in enumerate(reader.pages):
                txt = page.extract_text()
                if txt:
                    pages_text.append(f"--- Page {i+1} ---\n{txt}")
            text_content = "\n\n".join(pages_text)
            
        # 2. DOCX Documents (Extracts both paragraphs and table text)
        elif 'wordprocessingml.document' in ft or file_path_lower.endswith('.docx'):
            doc = docx.Document(file_path)
            full_text = []
            
            # Paragraph text
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text)
            
            # Table text (very common in resume layouts, CVs, and spreadsheets)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        # De-duplicate adjacent identical cell texts due to merged cells
                        unique_cells = []
                        for cell_val in row_cells:
                            if not unique_cells or unique_cells[-1] != cell_val:
                                unique_cells.append(cell_val)
                        full_text.append(" | ".join(unique_cells))
            
            text_content = "\n".join(full_text)
            
        # 3. DOC Documents (Classic binary Word doc)
        elif 'msword' in ft or file_path_lower.endswith('.doc'):
            text_content = "[Binary .doc file loaded. Direct text extraction is limited. Please use docx or pdf for best results.]"
            
        # 4. Text-based files (TXT, MD, CSV, JSON, code files, etc.)
        elif (
            'text/' in ft 
            or 'json' in ft 
            or 'csv' in ft 
            or 'xml' in ft 
            or 'yaml' in ft 
            or file_path_lower.endswith((
                '.txt', '.md', '.markdown', '.json', '.html', '.htm', '.css', '.js', 
                '.ts', '.jsx', '.tsx', '.py', '.java', '.c', '.cpp', '.h', '.cs', 
                '.go', '.rs', '.php', '.rb', '.sh', '.bat', '.ps1', '.sql', '.csv', 
                '.tsv', '.xml', '.yaml', '.yml', '.ini', '.conf', '.log', '.env'
            ))
        ):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
                
        # 5. Default fallback: detect binary vs text files
        else:
            # Check for null bytes to avoid dumping binary garbage as text
            is_binary = False
            if os.path.exists(file_path):
                try:
                    with open(file_path, 'rb') as f:
                        chunk = f.read(1024)
                        is_binary = b'\0' in chunk
                except Exception:
                    pass
            
            if is_binary:
                text_content = f"[Binary file '{os.path.basename(file_path)}' detected. Direct text extraction is not supported.]"
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
                    
    except Exception as e:
        print(f"Error extracting text from file {file_path}: {e}")
        text_content = f"Error reading document context: {str(e)}"
    
    return text_content

def analyze_file_content(client, uploaded_file_obj):
    if not client:
        uploaded_file_obj.extracted_text = "AI client not ready."
        uploaded_file_obj.summary = "Cannot analyze without AI client config."
        uploaded_file_obj.save()
        return

    file_path = uploaded_file_obj.file.path
    file_type = uploaded_file_obj.file_type.lower()
    file_name = uploaded_file_obj.file_name

    # Handling Images (OCR & vision summary)
    if any(img_type in file_type for img_type in ['image/', 'jpeg', 'png', 'gif', 'tiff']):
        try:
            with open(file_path, 'rb') as f:
                image_bytes = f.read()

            prompt = (
                "Please analyze this user-uploaded image. Do two things in detail:\n"
                "1. Under a heading 'EXTRACTED TEXT / OCR', extract any text readable in the image verbatim.\n"
                "2. Under a heading 'IMAGE DETAILS', describe the contents, style, context, and key objects/elements of the image.\n"
                "Also, at the very end, output a JSON block wrapped in ```json ... ``` containing exactly:\n"
                "{\n"
                "  \"summary\": \"A short 1-sentence descriptor of what this image is.\",\n"
                "  \"topics\": [\"keyword1\", \"keyword2\", ...]\n"
                "}"
            )

            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[
                    types.Part.from_bytes(data=image_bytes, mime_type=file_type),
                    prompt
                ]
            )

            result_text = response.text or ""
            uploaded_file_obj.extracted_text = result_text
            
            # Try to parse the trailing JSON block for summary/topics
            try:
                if "```json" in result_text:
                    json_str = result_text.split("```json")[1].split("```")[0].strip()
                    parsed = json.loads(json_str)
                    uploaded_file_obj.summary = parsed.get("summary", f"Image: {file_name}")
                    uploaded_file_obj.topics = parsed.get("topics", ["Image", "Visual"])
                else:
                    uploaded_file_obj.summary = f"Visual analysis of image '{file_name}' done."
                    uploaded_file_obj.topics = ["Image", "OCR"]
            except Exception:
                uploaded_file_obj.summary = f"Visual analysis of image '{file_name}'."
                uploaded_file_obj.topics = ["Image"]

        except Exception as e:
            uploaded_file_obj.extracted_text = f"Failed image parsing: {str(e)}"
            uploaded_file_obj.summary = f"Image '{file_name}' processing error."
            uploaded_file_obj.topics = ["Error"]

    # Handling Documents (PDF, Docx, Text, Markdown)
    else:
        # Extract text first
        text = extract_text_from_file(file_path, file_type)
        uploaded_file_obj.extracted_text = text if text else "Empty text extracted."
        
        if len(text.strip()) > 0:
            try:
                text_preview = text[:20000]
                prompt = (
                    f"You are analyzing an uploaded document named: '{file_name}'.\n"
                    f"First 20,000 characters of text:\n{text_preview}\n\n"
                    "Generate a JSON structure summarizing the main aspects of this document. "
                    "Return ONLY valid JSON. Response must contain two fields: 'summary' (a brief 1-2 sentence overview of the file) and 'topics' (a list of 3-7 core topics or entities)."
                )

                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=[prompt],
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json"
                    )
                )

                response_data = response.text or "{}"
                parsed = json.loads(response_data)
                uploaded_file_obj.summary = parsed.get("summary", "Document compiled and extracted.")
                uploaded_file_obj.topics = parsed.get("topics", ["Document"])
            except Exception as e:
                print(f"Error building document summary: {e}")
                uploaded_file_obj.summary = "Text content extracted. Detailed AI summary unavailable due to timeout/error."
                uploaded_file_obj.topics = ["Document"]
        else:
            uploaded_file_obj.summary = "No text content detected in this file."
            uploaded_file_obj.topics = ["Empty File"]

    uploaded_file_obj.save()
